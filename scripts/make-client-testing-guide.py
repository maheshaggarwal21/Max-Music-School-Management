# -*- coding: utf-8 -*-
"""Generate the client-facing Word testing guide for the MaxMusic platform."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x5B, 0x8D, 0xEF)      # Steel Blue
DARK  = RGBColor(0x1F, 0x2A, 0x44)
GREY  = RGBColor(0x55, 0x5B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Base font
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK

def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=10.5, fill=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if fill:
        shade(cell, fill)

def H(text, level=1, color=BRAND, size=None):
    p = doc.add_heading('', level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def para(text=None, bold=False, italic=False, color=None, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or DARK
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def rule_table(rows, widths=None, header=True):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            if header and ri == 0:
                set_cell(cell, val, bold=True, color=WHITE, fill='5B8DEF')
            else:
                set_cell(cell, val)
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.cell(ri, ci).width = Inches(w)
    return t

# ───────────────────────── COVER ─────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('MaxMusic School Platform')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Client Testing Guide')
r.font.size = Pt(15)
r.font.color.rgb = GREY

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run('A white-label, multi-school music academy platform — Demo environment')
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = GREY
doc.add_paragraph()

# ───────────────────────── WELCOME ─────────────────────────
H('1.  Welcome — what you are testing', 1)
para('This platform lets one operator (you) run many independent music schools from a single system. '
     'Each school gets its own fully-branded website with three sections — an Admin office, a Teacher '
     'area, and a Student area — while you keep a private master "Operator Console" behind the scenes.')
para('For this demo we have set up one sample school called ', space_after=0).add_run('“Demo Music School.”').bold = True
para('You will log into each of the four sections below using the credentials provided, and click around '
     'to see everything working live.')

# ───────────────────────── BEFORE YOU START ─────────────────────────
H('2.  Before you start — please read (60 seconds)', 1)

H('a)  The first page may take 30–60 seconds to open', 2, color=DARK, size=12)
para('Our demo server "sleeps" when nobody is using it, to save cost. The very first time you open a link '
     'it has to wake up, which can take up to a minute. Just wait, or refresh once. After that, everything '
     'is fast. This delay does not happen on a paid production server.')

H('b)  Every screen offers two ways to log in', 2, color=DARK, size=12)
bullet('  — sign in with the email/mobile and password we provide.', bold_lead='Password')
bullet('  — sign in with a phone number and a one-time code (like most modern apps).', bold_lead='OTP')
para('On each login screen you will see a ', space_after=0).add_run('Password').bold = True
para('Both work. Use whichever you like — credentials for both are in the table on the next page.')

H('c)  The demo OTP code', 2, color=DARK, size=12)
p = para('Because this is a demo, we are not sending real text messages yet. ')
p.add_run('When a screen asks for an OTP code, simply type: ')
r = p.add_run('  88990011')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = BRAND
para('(In the live product, each person receives their own unique code by SMS on their phone.)', italic=True, color=GREY)

H('d)  Why the school pages say “Demo Music School”, not “MaxMusic”', 2, color=DARK, size=12)
para('This is intentional and is the core idea of the product. Every school is "white-labelled" — it shows '
     'only its own name, logo and colour. Students and teachers never see the MaxMusic brand. Only you, in '
     'the private Operator Console, see the whole platform.')

doc.add_page_break()

# ───────────────────────── PANELS AT A GLANCE ─────────────────────────
H('3.  The four sections at a glance', 1)
rule_table([
    ['Section', 'Who uses it', 'In short'],
    ['Operator Console', 'You (the platform owner)', 'Master control room over every school. Private.'],
    ['Admin Panel', 'The school owner', 'Runs one school: students, teachers, batches, fees.'],
    ['Teacher Panel', 'Each teacher', 'Their classes, attendance, students.'],
    ['Student Panel', 'Each student', 'Their schedule, attendance, payments, practice tools.'],
], widths=[1.7, 2.0, 3.3])
doc.add_paragraph()

# ───────────────────────── MASTER LOGIN TABLE ─────────────────────────
H('4.  Login credentials & links (quick reference)', 1)
para('Tip: open each link in its own browser tab. Password login for the Operator and Admin uses an '
     'email; for Teacher and Student it uses a mobile number.', italic=True, color=GREY)

t = rule_table([
    ['Section / Link', 'Password login', 'OTP login'],
    ['OPERATOR CONSOLE\nmaxmusic-operator.vercel.app/login',
     'Email:  admin@maxmusic.internal\nPassword:  Operator@123',
     'Mobile:  9000000001\nCode:  88990011'],
    ['ADMIN PANEL\nmaxmusic-admin.vercel.app/demo-music-academy/admin/login',
     'Email:  teacher@demo.internal\nPassword:  Teacher@123',
     'Mobile:  9999999999\nCode:  88990011'],
    ['TEACHER PANEL\nmaxmusic-teacher.vercel.app/demo-music-academy/teacher/login',
     'Mobile:  7777777777\nPassword:  Teacher@123',
     'Mobile:  7777777777\nCode:  88990011'],
    ['STUDENT PANEL\nmaxmusic-student.vercel.app/demo-music-academy/student/login',
     'Mobile:  6666666666\nPassword:  Student@123',
     'Mobile:  6666666666\nCode:  88990011'],
], widths=[2.9, 2.1, 2.0])
doc.add_paragraph()

para('How to log in with OTP (same steps on every screen):', bold=True)
doc.add_paragraph('Click the “OTP” tab on the login box.', style='List Number')
doc.add_paragraph('Type the mobile number shown above for that section.', style='List Number')
doc.add_paragraph('Click “Send OTP” / “Get code”.', style='List Number')
doc.add_paragraph('Enter the demo code 88990011 and press Sign in.', style='List Number')

doc.add_page_break()

# ───────────────────────── DEEP DIVE ─────────────────────────
H('5.  What to explore in each section', 1)

def login_line(pw, otp):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Log in:  '); r.bold = True; r.font.color.rgb = BRAND
    p.add_run('Password — ' + pw + '   •   OTP — ' + otp)

# Operator
H('A.  Operator Console  (your master control room)', 2, color=DARK, size=13)
para('Link:  https://maxmusic-operator.vercel.app/login', bold=True, color=GREY)
login_line('admin@maxmusic.internal / Operator@123', 'mobile 9000000001, code 88990011')
para('This is the only screen that shows the whole platform. Students and teachers can never reach it.')
bullet('platform-wide totals — schools, students, teachers, fees collected.', bold_lead='Dashboard: ')
bullet('every school on the platform. Open one to see its full detail. Use “Impersonate” to step inside that school’s Admin panel, and “Grant / Revoke Admin” to switch a teacher between salaried and independent.', bold_lead='Institutions: ')
bullet('approve requests from schools that want to change their web address.', bold_lead='Slug Requests: ')
bullet('see and edit every student / teacher across all schools (god view).', bold_lead='Students / Teachers: ')
bullet('reset anyone’s forgotten password — the system issues a one-time temporary password.', bold_lead='Credentials: ')
bullet('a complete, tamper-proof history of every action — who did what, and when.', bold_lead='Changes History: ')
bullet('platform defaults (rent), the instrument catalogue, and the master demo OTP code.', bold_lead='Settings: ')
doc.add_paragraph()

# Admin
H('B.  Admin Panel  (the school owner’s office)', 2, color=DARK, size=13)
para('Link:  https://maxmusic-admin.vercel.app/demo-music-academy/admin/login', bold=True, color=GREY)
login_line('teacher@demo.internal / Teacher@123', 'mobile 9999999999, code 88990011')
para('Branded entirely as “Demo Music School”. This is what a real school owner sees every day.')
bullet('this school’s numbers, plus one-click shortcuts.', bold_lead='Dashboard: ')
bullet('approve a pending enrollment — it becomes a real student with a temporary password.', bold_lead='New Requests: ')
bullet('manage the people in this school.', bold_lead='Students / Teachers: ')
bullet('a class group = instrument + days + time + teacher.', bold_lead='Batches: ')
bullet('fee-and-duration templates (e.g. “Beginner — 12 classes”).', bold_lead='Class levels: ')
bullet('reusable day-patterns and time-windows that combine into batches.', bold_lead='Suitable Days / Times: ')
bullet('mark and correct attendance for any batch.', bold_lead='Attendance: ')
bullet('record fees received and reconcile payments.', bold_lead='Payments: ')
bullet('change the school’s name, logo and colour, or request a new web address.', bold_lead='Settings: ')
doc.add_paragraph()

# Teacher
H('C.  Teacher Panel', 2, color=DARK, size=13)
para('Link:  https://maxmusic-teacher.vercel.app/demo-music-academy/teacher/login', bold=True, color=GREY)
login_line('mobile 7777777777 / Teacher@123', 'mobile 7777777777, code 88990011')
para('We use the Guitar teacher here because they have a live batch with students and attendance to show.')
bullet('today’s classes and any declared holidays.', bold_lead='Dashboard: ')
bullet('your class groups. Open one to launch a live session, mark attendance, view students, or archive it.', bold_lead='My Batches: ')
bullet('one-tap Present / Absent — it saves instantly and updates the student’s record.', bold_lead='Attendance: ')
bullet('the teacher roster with simple performance indicators.', bold_lead='Teachers: ')
bullet('your own details — try editing and saving.', bold_lead='Profile: ')
doc.add_paragraph()

# Student
H('D.  Student Panel', 2, color=DARK, size=13)
para('Link:  https://maxmusic-student.vercel.app/demo-music-academy/student/login', bold=True, color=GREY)
login_line('mobile 6666666666 / Student@123', 'mobile 6666666666, code 88990011')
para('This student is enrolled in a guitar batch, so you will see real classes, a timetable and attendance.')
bullet('next class, attendance %, plan validity, and the “My Plan” summary.', bold_lead='Dashboard: ')
bullet('the student’s class history.', bold_lead='My Classes: ')
bullet('a clean weekly schedule grid.', bold_lead='Timetable: ')
bullet('fee history.', bold_lead='Payments: ')
bullet('a built-in Metronome and a Guitar Chords library to practise with.', bold_lead='Practice Tools: ')
bullet('the student’s teacher and the school’s contact details (no MaxMusic branding).', bold_lead='Contact Us: ')

doc.add_page_break()

# ───────────────────────── GUIDED TOUR ─────────────────────────
H('6.  Suggested 10-minute guided tour', 1)
para('Follow this path to see the whole system connect end-to-end:')
steps = [
    ('Start in the Operator Console.', 'Look at the platform totals, then open “Demo Music School” and click Impersonate — you are now inside that school’s Admin panel as its owner.'),
    ('In the Admin panel, open New Requests.', 'Approve a pending enrollment (or use Add Student). Notice it becomes a real student with a temporary password.'),
    ('Open Batches.', 'See how a class is built from an instrument, a day-pattern, a time-slot and a teacher.'),
    ('Log out and open the Teacher panel (7777777777).', 'Go to My Batches → open the batch → mark a student Present. It saves instantly.'),
    ('Log out and open the Student panel (6666666666).', 'See the attendance you just marked reflected here, check the Timetable, then try the Metronome under Practice Tools.'),
]
for i, (lead, rest) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(lead + '  '); r.bold = True
    p.add_run(rest)
para()
para('That single loop — operator → school → teacher → student — is the heart of the platform.', italic=True, color=BRAND)

# ───────────────────────── FAQ ─────────────────────────
H('7.  Questions you might have', 1)
faq = [
    ('The page is taking a long time to load.',
     'The free demo server was asleep and is waking up (up to ~60 seconds). Refresh once; after that it is fast. This does not happen on a production server.'),
    ('I didn’t receive an SMS for the OTP.',
     'The demo doesn’t send real texts yet. Use the demo code 88990011. In the live product every user gets their own code by SMS.'),
    ('Why do the school pages not mention MaxMusic?',
     'By design. Each school is white-labelled as itself. Students must never see the operator brand. Only the Operator Console shows the whole platform.'),
    ('Will my clicking around break anything?',
     'No. This is a demo database with sample data — explore freely. Changes you make (adding a student, marking attendance) are real within the demo and show up across the panels, which is exactly how you can follow the flow.'),
    ('A screen sent me back to the login page.',
     'Just log in again — demo sessions are short-lived for security.'),
]
for q, a in faq:
    p = para(space_after=2); r = p.add_run('Q.  ' + q); r.bold = True; r.font.color.rgb = DARK
    para('A.  ' + a, space_after=8)

# ───────────────────────── SECURITY NOTE ─────────────────────────
H('8.  Note for go-live', 1)
bullet('The demo code 88990011 is a master test code. Before launch, change it in Operator → Settings and switch on real SMS so each person gets their own unique code.', bold_lead='Change the demo code: ')
bullet('All passwords above are demo passwords — set fresh ones for real users.', bold_lead='Reset passwords: ')
bullet('On a paid server the wake-up delay disappears and the platform runs at full speed.', bold_lead='Production speed: ')

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run('MaxMusic School Platform — Demo testing guide. For questions, contact your development team.')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

out = r'c:\Users\Mahesh\Desktop\MaxMusic-Client-Testing-Guide.docx'
doc.save(out)
print('SAVED:', out)
